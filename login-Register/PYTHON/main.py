# from fastapi import FastAPI, UploadFile, File, HTTPException
# from fastapi.middleware.cors import CORSMiddleware
# from transformers import pipeline, AutoTokenizer, AutoModelForSeq2SeqLM
# import pdfplumber
# from docx import Document
# import os

# app = FastAPI()

# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=["*"],
#     allow_credentials=True,
#     allow_methods=["*"],
#     allow_headers=["*"],
# )

# tokenizer = AutoTokenizer.from_pretrained("facebook/bart-large-cnn")
# model = AutoModelForSeq2SeqLM.from_pretrained("facebook/bart-large-cnn")
# summarizer = pipeline("summarization", model=model, tokenizer=tokenizer)

# def read_pdf(file_path):
#     text = ""
#     try:
#         with pdfplumber.open(file_path) as pdf:
#             for page in pdf.pages:
#                 text += page.extract_text() + "\n"
#     except Exception as e:
#         print(f"Error reading PDF: {e}")
#     return text

# def read_docx(file_path):
#     text = ""
#     try:
#         doc = Document(file_path)
#         for para in doc.paragraphs:
#             text += para.text + "\n"
#     except Exception as e:
#         print(f"Error reading DOCX: {e}")
#     return text

# def summarize_text(text, max_length=500, min_length=150):
#     inputs = tokenizer(text, max_length=1024, truncation=True, return_tensors="pt")
#     summary_ids = model.generate(inputs["input_ids"], num_beams=4, max_length=max_length, min_length=min_length)
#     summary = tokenizer.decode(summary_ids[0], skip_special_tokens=True)
#     return summary

# @app.post("/summarize/")
# async def summarize(file: UploadFile = File(None), text: str = None):
#     try:
#         if file:
#             file_path = f"uploaded_docs/{file.filename}"
#             with open(file_path, "wb") as buffer:
#                 buffer.write(file.file.read())
            
#             if file.content_type == "application/pdf":
#                 text = read_pdf(file_path)
#             elif file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#                 text = read_docx(file_path)
#             else:
#                 os.remove(file_path)
#                 raise HTTPException(status_code=400, detail="Unsupported file type")

#             os.remove(file_path)

#         if not text:
#             raise HTTPException(status_code=400, detail="No text provided for summarization")

#         summary = summarize_text(text)
#         return {"summary": summary}
#     except Exception as e:
#         print(f"Internal Server Error: {e}")
#         raise HTTPException(status_code=500, detail="Internal Server Error")

# if __name__ == "__main__":
#     import uvicorn
#     uvicorn.run(app, host="0.0.0.0", port=8000)


# from fastapi import FastAPI, UploadFile, File, HTTPException
# from fastapi.middleware.cors import CORSMiddleware
# from transformers import pipeline, AutoTokenizer, AutoModelForSeq2SeqLM
# import pdfplumber
# from docx import Document
# import os

# app = FastAPI()

# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=["*"],
#     allow_credentials=True,
#     allow_methods=["*"],
#     allow_headers=["*"],
# )

# # Load tokenizer and model from local directory
# model_directory = "../model/bart-large-cnn"
# tokenizer = AutoTokenizer.from_pretrained(model_directory)
# model = AutoModelForSeq2SeqLM.from_pretrained(model_directory)
# summarizer = pipeline("summarization", model=model, tokenizer=tokenizer)

# def read_pdf(file_path):
#     text = ""
#     try:
#         with pdfplumber.open(file_path) as pdf:
#             for page in pdf.pages:
#                 text += page.extract_text() + "\n"
#     except Exception as e:
#         print(f"Error reading PDF: {e}")
#     return text

# def read_docx(file_path):
#     text = ""
#     try:
#         doc = Document(file_path)
#         for para in doc.paragraphs:
#             text += para.text + "\n"
#     except Exception as e:
#         print(f"Error reading DOCX: {e}")
#     return text

# def summarize_text(text, max_length=500, min_length=150):
#     inputs = tokenizer(text, max_length=1024, truncation=True, return_tensors="pt")
#     summary_ids = model.generate(inputs["input_ids"], num_beams=4, max_length=max_length, min_length=min_length)
#     summary = tokenizer.decode(summary_ids[0], skip_special_tokens=True)
#     return summary

# @app.post("/summarize/")
# async def summarize(file: UploadFile = File(None), text: str = None):
#     try:
#         if file:
#             file_path = f"uploaded_docs/{file.filename}"
#             with open(file_path, "wb") as buffer:
#                 buffer.write(file.file.read())
            
#             if file.content_type == "application/pdf":
#                 text = read_pdf(file_path)
#             elif file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#                 text = read_docx(file_path)
#             else:
#                 os.remove(file_path)
#                 raise HTTPException(status_code=400, detail="Unsupported file type")

#             os.remove(file_path)

#         if not text:
#             raise HTTPException(status_code=400, detail="No text provided for summarization")

#         summary = summarize_text(text)
#         return {"summary": summary}
#     except Exception as e:
#         print(f"Internal Server Error: {e}")
#         raise HTTPException(status_code=500, detail="Internal Server Error")

# if __name__ == "_main_":
#     import uvicorn
#     uvicorn.run(app, host="0.0.0.0", port=8000)

from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from transformers import pipeline, AutoTokenizer, AutoModelForSeq2SeqLM
import pdfplumber
from docx import Document
import os

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Load tokenizer and model from local directory
model_directory = "../model/bart-large-cnn"
tokenizer = AutoTokenizer.from_pretrained(model_directory)
model = AutoModelForSeq2SeqLM.from_pretrained(model_directory)
summarizer = pipeline("summarization", model=model, tokenizer=tokenizer)

def read_pdf(file_path):
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() + "\n"
    except Exception as e:
        print(f"Error reading PDF: {e}")
    return text

def read_docx(file_path):
    text = ""
    try:
        doc = Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
    except Exception as e:
        print(f"Error reading DOCX: {e}")
    return text

def chunk_text(text, chunk_size=3000):
    """Split the text into chunks of the given size."""
    return [text[i:i + chunk_size] for i in range(0, len(text), chunk_size)]

def summarize_text(text, max_length=1000, min_length=150):
    inputs = tokenizer(text, max_length=1024, truncation=True, return_tensors="pt")
    summary_ids = model.generate(inputs["input_ids"], num_beams=4, max_length=max_length, min_length=min_length)
    summary = tokenizer.decode(summary_ids[0], skip_special_tokens=True)
    return summary

@app.post("/summarize/")
async def summarize(file: UploadFile = File(None), text: str = None):
    try:
        if file:
            file_path = f"uploaded_docs/{file.filename}"
            with open(file_path, "wb") as buffer:
                buffer.write(file.file.read())
            
            if file.content_type == "application/pdf":
                text = read_pdf(file_path)
            elif file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                text = read_docx(file_path)
            else:
                os.remove(file_path)
                raise HTTPException(status_code=400, detail="Unsupported file type")

            os.remove(file_path)

        if not text:
            raise HTTPException(status_code=400, detail="No text provided for summarization")

        # Break the text into chunks and summarize each chunk
        chunks = chunk_text(text)
        summaries = [summarize_text(chunk) for chunk in chunks]

        # Summarize the combined summaries to get the final summary
        final_summary = summarize_text(" ".join(summaries))

        return {"summary": final_summary}
    except Exception as e:
        print(f"Internal Server Error: {e}")
        raise HTTPException(status_code=500, detail="Internal Server Error")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
